import pandas as pd
from rapidfuzz import fuzz, process
from docx import Document
from docx.shared import RGBColor
from io import BytesIO
from pydocx import PyDocX  # .doc ファイルを扱うためのライブラリ
import streamlit as st
from PyPDF2 import PdfReader  # PDFからテキストを抽出するためのライブラリ

# Excel ファイルを読み込む関数
def load_excel(file):
    # Handle dynamic column names and ensure proper format
    df = pd.read_excel(file, engine="openpyxl")
    if df.columns.size < 1:
        raise ValueError("The Excel file must contain at least one column with terms.")
    return df

# Word、DOC または PDF ファイルからテキストを抽出する関数
def extract_text_from_file(file, file_type):
    if file_type == "docx":
        doc = Document(file)
        return "\n".join([paragraph.text for paragraph in doc.paragraphs])
    elif file_type == "doc":
        return PyDocX.to_text(file)
    elif file_type == "pdf":
        reader = PdfReader(file)
        text = ""
        for page in reader.pages:
            page_text = page.extract_text()
            # 改行と複数スペースを削除
            page_text = page_text.replace("\n", " ").replace("\r", " ")
            page_text = " ".join(page_text.split())
            text += page_text + " "
        # 全体をさらに正規化
        text = text.strip()
        return text
    else:
        return ""

# Fuzzy Matching を用いて類似語を検出する関数
def find_similar_terms_with_context(text, terms, threshold):
    detected_terms = []

    words = text.split()
    for word in words:
        matches = process.extract(word, terms, scorer=fuzz.partial_ratio, limit=10)
        for match in matches:
            if match[1] >= threshold and match[1] < 100:  # Include matches above the threshold but exclude exact matches
                start_index = text.find(word)
                context = text[max(0, start_index-10):start_index+10+len(word)]
                detected_terms.append((word, match[0], match[1], context, start_index))

    return detected_terms

# 修正を適用して新しい Word ファイルを作成する関数
def create_corrected_word_file_with_formatting(original_text, corrections):
    doc = Document()
    for paragraph_text in original_text.split("\n"):
        paragraph = doc.add_paragraph()
        start_index = 0

        # Apply corrections with yellow highlighting
        for incorrect, correct in corrections:
            while incorrect in paragraph_text[start_index:]:
                start_index = paragraph_text.find(incorrect, start_index)
                end_index = start_index + len(incorrect)

                # Add text before the match
                paragraph.add_run(paragraph_text[:start_index])

                # Add the corrected text with yellow highlighting
                run = paragraph.add_run(correct)
                run.font.highlight_color = 6  # Yellow highlighting

                # Update the remaining text
                paragraph_text = paragraph_text[end_index:]
                start_index = 0

        # Add any remaining text
        paragraph.add_run(paragraph_text)

    output = BytesIO()
    doc.save(output)
    output.seek(0)
    return output

# 修正を適用して新しい Word ファイルを作成する関数
def create_correction_table_with_context(detected):
    correction_table = pd.DataFrame(detected, columns=["原稿内の語", "類似する用語", "類似度", "文脈", "位置"])
    return correction_table

# 正誤表を使用して修正を適用する関数
def apply_corrections_with_table(text, correction_df):
    corrections = []
    for _, row in correction_df.iterrows():
        incorrect, correct = row.iloc[0], row.iloc[1]  # Use the first column as "incorrect" and the second as "correct"
        if incorrect in text:
            corrections.append((incorrect, correct))
        text = text.replace(incorrect, correct)
    return text, corrections

# 利用漢字表を使用して修正を適用する関数
def apply_kanji_table(text, kanji_df):
    corrections = []
    for _, row in kanji_df.iterrows():
        hiragana, kanji = row.iloc[0], row.iloc[1]  # Use the first column as "hiragana" and the second as "kanji"
        if hiragana in text:
            corrections.append((hiragana, kanji))
        text = text.replace(hiragana, kanji)
    return text, corrections

# Streamlit アプリケーション
st.markdown("<h1 style='text-align: center;'>南江堂様用用語チェックサービス（笑）</h1>", unsafe_allow_html=True)

st.write("以下のファイルを個別にアップロードしてください（正誤表と利用漢字表は実質的には同じことをするものです）:")
word_file = st.file_uploader("原稿ファイル (Word, DOC, PDF):", type=["docx", "doc", "pdf"])
terms_file = st.file_uploader("用語集ファイル (A列に正しい用語を入れたExcel。A列以外には何も入れないでください):", type=["xlsx"])
correction_file = st.file_uploader("正誤表ファイル (A列に誤った語句をB列にそれに対する正しい語句を入れたExcel。AとB列以外には何も入れないでください):", type=["xlsx"])
kanji_file = st.file_uploader("利用漢字表ファイル (A列にひらがなを、B列に漢字を入れたExcel。AとB列以外には何も入れないでください):", type=["xlsx"])

if word_file and (terms_file or correction_file or kanji_file):
    # 原稿ファイルの読み込み
    file_type = word_file.name.split(".")[-1]
    original_text = extract_text_from_file(word_file, file_type)

    corrections = []

    # 用語集がアップロードされている場合
    if terms_file:
        try:
            terms_df = load_excel(terms_file)
            terms = terms_df.iloc[:, 0].dropna().astype(str).tolist()

            # 類似度の閾値を入力
            threshold = st.slider("類似度の閾値を設定してください (50-99):", min_value=50, max_value=99, value=65)
            detected = find_similar_terms_with_context(original_text, terms, threshold)

            # 結果を表示
            if detected:
                st.success("類似語が検出されました！")
                correction_table = create_correction_table_with_context(detected)
                st.dataframe(correction_table)

                # 修正箇所を表形式でダウンロード
                output = BytesIO()
                with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                    correction_table.to_excel(writer, index=False, sheet_name="Corrections")
                st.download_button(
                    label="修正箇所をダウンロード",
                    data=output.getvalue(),
                    file_name="correction_table_with_context.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

            else:
                st.warning("類似する語は検出されませんでした。")

        except Exception as e:
            st.error(f"用語集ファイルの読み込み中にエラーが発生しました: {e}")

    # 正誤表がアップロードされている場合
    if correction_file:
        try:
            correction_df = load_excel(correction_file)
            original_text, corrections_from_table = apply_corrections_with_table(original_text, correction_df)
            corrections.extend(corrections_from_table)
            st.success("正誤表を適用しました！")
            if corrections_from_table:
                st.write("修正内容:")
                corrections_df = pd.DataFrame(corrections_from_table, columns=["誤った用語", "正しい用語"])
                st.dataframe(corrections_df)

                # 修正箇所をダウンロード
                output = BytesIO()
                with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                    corrections_df.to_excel(writer, index=False, sheet_name="Corrections")
                st.download_button(
                    label="正誤表修正箇所をダウンロード",
                    data=output.getvalue(),
                    file_name="corrections_from_table.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

        except Exception as e:
            st.error(f"正誤表の処理中にエラーが発生しました: {e}")

    # 利用漢字表がアップロードされている場合
    if kanji_file:
        try:
            kanji_df = load_excel(kanji_file)
            original_text, kanji_corrections = apply_kanji_table(original_text, kanji_df)
            corrections.extend(kanji_corrections)
            st.success("利用漢字表を適用しました！")
            if kanji_corrections:
                st.write("修正内容:")
                kanji_corrections_df = pd.DataFrame(kanji_corrections, columns=["ひらがな", "漢字"])
                st.dataframe(kanji_corrections_df)

                # 修正箇所をダウンロード
                output = BytesIO()
                with pd.ExcelWriter(output, engine="xlsxwriter") as writer:
                    kanji_corrections_df.to_excel(writer, index=False, sheet_name="Kanji Corrections")
                st.download_button(
                    label="利用漢字表修正箇所をダウンロード",
                    data=output.getvalue(),
                    file_name="kanji_corrections.xlsx",
                    mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet",
                )

        except Exception as e:
            st.error(f"利用漢字表の処理中にエラーが発生しました: {e}")

else:
    st.warning("原稿ファイルと、用語集、正誤表、利用漢字表のいずれかをアップロードしてください！")
